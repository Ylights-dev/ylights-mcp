#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Ylights Tools MCP Server v3
Персональная библиотека AI-инструментов Владимира.
Промпты не хардкодятся — подтягиваются git pull из репозитория Yl_Promts
(каждый .md файл с YAML frontmatter = один инструмент).
"""
import asyncio
import io
import json
import os
import subprocess
import sys
import time
from contextlib import asynccontextmanager

from mcp.server import Server
from mcp.server.sse import SseServerTransport
from mcp.types import Tool, TextContent, Prompt, PromptMessage
from starlette.applications import Starlette
from starlette.middleware.base import BaseHTTPMiddleware
from starlette.requests import Request
from starlette.responses import FileResponse, JSONResponse, PlainTextResponse
from starlette.routing import Route, Mount
from starlette.staticfiles import StaticFiles
import uvicorn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

STATIC_DIR = os.path.join(os.path.dirname(__file__), "static")
PROMPTS_REPO_URL = os.environ.get("PROMPTS_REPO_URL", "https://github.com/Ylights-dev/Yl_Promts.git")
PROMPTS_REPO_DIR = os.environ.get("PROMPTS_REPO_DIR", "/data/Yl_Promts")
PROMPTS_SYNC_INTERVAL = int(os.environ.get("PROMPTS_SYNC_INTERVAL", "1800"))  # 30 мин
MCP_API_KEY = os.environ.get("MCP_API_KEY", "")

server = Server("ylights-tools")

# ──────────────────────────────────────────────
# Git-sync библиотеки промптов
# ──────────────────────────────────────────────

TOOLS: list[Tool] = []
PROMPT_TEMPLATES: dict[str, str] = {}
TOOL_CATEGORY: dict[str, str] = {}
TOOL_DEFAULTS: dict[str, dict] = {}
_last_sync_error = None
_last_sync_time = None


def _parse_frontmatter(text: str):
    if not text.startswith("---"):
        return {}, text
    end = text.find("\n---", 3)
    if end == -1:
        return {}, text
    header = text[3:end].strip()
    body = text[end + 4:].lstrip("\n")
    meta = {"inputs": []}
    current_input = None
    for line in header.splitlines():
        if not line.strip():
            continue
        if line.startswith("  - name:"):
            current_input = {"name": line.split(":", 1)[1].strip()}
            meta["inputs"].append(current_input)
        elif line.startswith("    description:") and current_input is not None:
            current_input["description"] = line.split(":", 1)[1].strip()
        elif line.startswith("    required:") and current_input is not None:
            current_input["required"] = line.split(":", 1)[1].strip().lower() == "true"
        elif line.startswith("    default:") and current_input is not None:
            v = line.split(":", 1)[1].strip()
            current_input["default"] = v.strip('"')
        elif line.startswith("tool:"):
            meta["tool"] = line.split(":", 1)[1].strip()
        elif line.startswith("description:"):
            meta["description"] = line.split(":", 1)[1].strip()
    return meta, body


def _load_prompts_from_repo():
    """Сканирует PROMPTS_REPO_DIR и строит TOOLS/PROMPT_TEMPLATES."""
    global TOOLS, PROMPT_TEMPLATES, TOOL_CATEGORY, TOOL_DEFAULTS
    tools, templates, categories, defaults = [], {}, {}, {}

    if not os.path.isdir(PROMPTS_REPO_DIR):
        return tools, templates, categories, defaults

    for category in sorted(os.listdir(PROMPTS_REPO_DIR)):
        cat_path = os.path.join(PROMPTS_REPO_DIR, category)
        if not os.path.isdir(cat_path) or category.startswith("."):
            continue
        for fname in sorted(os.listdir(cat_path)):
            if not fname.endswith(".md"):
                continue
            fpath = os.path.join(cat_path, fname)
            try:
                with open(fpath, "r", encoding="utf-8") as f:
                    raw = f.read()
            except OSError:
                continue
            meta, body = _parse_frontmatter(raw)
            name = meta.get("tool") or os.path.splitext(fname)[0]
            description = meta.get("description", name)
            props, required_list, tool_defaults = {}, [], {}
            for inp in meta.get("inputs", []):
                iname = inp.get("name")
                if not iname:
                    continue
                prop = {"type": "string", "description": inp.get("description", "")}
                if "default" in inp:
                    prop["default"] = inp["default"]
                    tool_defaults[iname] = inp["default"]
                props[iname] = prop
                if inp.get("required"):
                    required_list.append(iname)
            tools.append(Tool(
                name=name,
                description=f"[{category}] {description}",
                inputSchema={"type": "object", "properties": props, "required": required_list},
            ))
            templates[name] = body.strip()
            categories[name] = category
            defaults[name] = tool_defaults

    return tools, templates, categories, defaults


def _git_sync():
    """git clone (первый запуск) или git pull (обновление) репозитория промптов."""
    global _last_sync_error, _last_sync_time
    try:
        if not os.path.isdir(os.path.join(PROMPTS_REPO_DIR, ".git")):
            os.makedirs(os.path.dirname(PROMPTS_REPO_DIR) or ".", exist_ok=True)
            subprocess.run(["git", "clone", "--depth", "1", PROMPTS_REPO_URL, PROMPTS_REPO_DIR],
                            check=True, capture_output=True, text=True, timeout=60)
        else:
            subprocess.run(["git", "-C", PROMPTS_REPO_DIR, "pull", "--ff-only"],
                            check=True, capture_output=True, text=True, timeout=60)
        _last_sync_error = None
    except subprocess.CalledProcessError as e:
        _last_sync_error = e.stderr or str(e)
    except Exception as e:
        _last_sync_error = str(e)
    _last_sync_time = time.time()

    global TOOLS, PROMPT_TEMPLATES, TOOL_CATEGORY, TOOL_DEFAULTS
    TOOLS, PROMPT_TEMPLATES, TOOL_CATEGORY, TOOL_DEFAULTS = _load_prompts_from_repo()


async def _auto_sync_loop():
    while True:
        await asyncio.sleep(PROMPTS_SYNC_INTERVAL)
        await asyncio.get_event_loop().run_in_executor(None, _git_sync)


# ──────────────────────────────────────────────
# Спец-инструменты (не из промпт-репозитория)
# ──────────────────────────────────────────────

SPECIAL_TOOLS = [
    Tool(name="docx_generator",
         description="[Инструменты] Генератор DOCX-документов из Markdown — для моделей, которые не умеют сами создавать файлы",
         inputSchema={"type": "object", "required": ["md_content"],
                      "properties": {
                          "md_content": {"type": "string", "description": "Markdown текст"},
                          "filename": {"type": "string", "description": "Имя файла", "default": "document"},
                      }}),
]


async def _generate_docx(args: dict) -> list[TextContent]:
    try:
        from docx import Document
        from docx.shared import Pt

        md = args["md_content"]
        fname = args.get("filename", "document")

        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        for line in md.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("## "):
                p = doc.add_paragraph()
                r = p.add_run(line[3:]); r.bold = True; r.font.size = Pt(14); r.font.name = "Calibri"
            elif line.startswith("### "):
                p = doc.add_paragraph()
                r = p.add_run(line[4:]); r.bold = True; r.font.size = Pt(12); r.font.name = "Calibri"
            elif line.startswith("# "):
                p = doc.add_paragraph()
                r = p.add_run(line[2:]); r.bold = True; r.font.size = Pt(16); r.font.name = "Calibri"
                p.alignment = 1
            elif line.startswith("- ") or line.startswith("* "):
                p = doc.add_paragraph(style="List Bullet")
                r = p.add_run(line[2:]); r.font.name = "Calibri"; r.font.size = Pt(11)
            elif line.startswith("**") and line.endswith("**"):
                p = doc.add_paragraph()
                r = p.add_run(line[2:-2]); r.bold = True; r.font.name = "Calibri"; r.font.size = Pt(11)
            else:
                p = doc.add_paragraph()
                r = p.add_run(line); r.font.name = "Calibri"; r.font.size = Pt(11)

        buf = io.BytesIO()
        doc.save(buf); buf.seek(0)
        import base64
        b64 = base64.b64encode(buf.read()).decode("utf-8")
        return [TextContent(type="text", text=f"Файл: {fname}.docx\nРазмер: {len(b64)//1024} KB\nBase64:\n{b64}")]
    except Exception as e:
        return [TextContent(type="text", text=f"Ошибка: {e}")]


# ──────────────────────────────────────────────
# MCP handlers
# ──────────────────────────────────────────────

@server.list_tools()
async def list_tools() -> list[Tool]:
    return TOOLS + SPECIAL_TOOLS


def _render_template(name: str, args: dict) -> str:
    template = PROMPT_TEMPLATES.get(name)
    if not template:
        text = f"## {name}\n\n### Входные данные:\n"
        for k, v in args.items():
            text += f"- {k}: {v}\n"
        return text + "\n### Сгенерируй ответ на основе введённых данных.\n"

    merged = dict(TOOL_DEFAULTS.get(name, {}))
    merged.update(args)
    try:
        return template.format(**merged)
    except KeyError as e:
        return template + f"\n\n[!] Не хватает параметра: {e}"


@server.call_tool()
async def call_tool(name: str, arguments: dict) -> list[TextContent]:
    if name == "docx_generator":
        return await _generate_docx(arguments)
    return [TextContent(type="text", text=_render_template(name, arguments))]


@server.list_prompts()
async def list_prompts() -> list[Prompt]:
    return []


@server.get_prompt()
async def get_prompt(name: str, arguments: dict | None = None) -> list[PromptMessage]:
    return []


# ──────────────────────────────────────────────
# HTTP / REST
# ──────────────────────────────────────────────

sse = SseServerTransport("/messages/")


class ApiKeyMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request: Request, call_next):
        if MCP_API_KEY and request.url.path.startswith("/api/"):
            key = request.headers.get("X-API-Key") or (
                request.headers.get("Authorization", "").removeprefix("Bearer ").strip()
            )
            if key != MCP_API_KEY:
                return JSONResponse({"error": "unauthorized"}, status_code=401)
        return await call_next(request)


async def handle_sse(request):
    async with sse.connect_sse(request.scope, request.receive, request._send) as streams:
        await server.run(streams[0], streams[1], server.create_initialization_options())


async def handle_messages(request):
    await sse.handle_post_message(request.scope, request.receive, request._send)


async def handle_root(request):
    index_path = os.path.join(STATIC_DIR, "index.html")
    if os.path.exists(index_path):
        return FileResponse(index_path)
    return PlainTextResponse("static/index.html not found", status_code=404)


async def api_tools(request):
    items = []
    for t in TOOLS + SPECIAL_TOOLS:
        items.append({
            "name": t.name,
            "description": t.description,
            "category": TOOL_CATEGORY.get(t.name, "Инструменты"),
            "inputSchema": t.inputSchema,
        })
    return JSONResponse({"tools": items})


async def api_call(request):
    body = await request.json()
    name = body.get("tool")
    args = body.get("arguments", {})
    if not name:
        return JSONResponse({"error": "tool is required"}, status_code=400)
    result = await call_tool(name, args)
    text = "\n".join(c.text for c in result if hasattr(c, "text"))
    return JSONResponse({"result": text})


async def api_sync(request):
    await asyncio.get_event_loop().run_in_executor(None, _git_sync)
    return JSONResponse({"tools_loaded": len(TOOLS), "error": _last_sync_error})


async def health(request):
    return JSONResponse({"status": "ok"})


async def api_status(request):
    return JSONResponse({
        "tools_loaded": len(TOOLS) + len(SPECIAL_TOOLS),
        "last_sync_error": _last_sync_error,
        "last_sync_time": _last_sync_time,
        "repo_dir": PROMPTS_REPO_DIR,
    })


@asynccontextmanager
async def lifespan(app):
    await asyncio.get_event_loop().run_in_executor(None, _git_sync)
    task = asyncio.create_task(_auto_sync_loop())
    yield
    task.cancel()


app = Starlette(routes=[
    Route("/", endpoint=handle_root),
    Route("/health", endpoint=health),
    Route("/sse", endpoint=handle_sse),
    Route("/messages/", endpoint=handle_messages, methods=["POST"]),
    Route("/api/tools", endpoint=api_tools, methods=["GET"]),
    Route("/api/call", endpoint=api_call, methods=["POST"]),
    Route("/api/sync", endpoint=api_sync, methods=["POST"]),
    Route("/api/status", endpoint=api_status, methods=["GET"]),
    Mount("/static", app=StaticFiles(directory=STATIC_DIR, html=False), name="static"),
], lifespan=lifespan)
app.add_middleware(ApiKeyMiddleware)


if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=3001)
